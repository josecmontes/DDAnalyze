"""
Phase 2 — Final Report Generator (Iterative Architecture)

Three-step process:
  Step A: Extract all successful analysis code → extracted_code.md
  Step B: LLM Architect plans the report, then LLM Writer drafts it section-by-section → final_report.md
  Step C: Convert the markdown report to a Word document → final_report.docx
          (embeds tables as Word tables, applies Deloitte colors, and inserts graph images)

Run this manually after loop.py finishes.
"""

import json
import logging
import os
import re
import sys
import time
from datetime import datetime
from pathlib import Path

import anthropic
import yaml
from dotenv import load_dotenv

# Change to the directory containing this script so relative paths work
os.chdir(Path(__file__).parent)
load_dotenv()
API_KEY = os.getenv('API_KEY')

# ─── Logging Setup ────────────────────────────────────────────────────────────

logger = logging.getLogger("ddanalyze.phase2")

def setup_logging(debug: bool = False, log_dir: str = "logs") -> Path:
    """Configure console + file logging. Returns the path of the log file."""
    Path(log_dir).mkdir(exist_ok=True)
    log_file = Path(log_dir) / f"phase2_{datetime.now().strftime('%Y%m%d_%H%M%S')}.log"

    console_level = logging.DEBUG if debug else logging.INFO
    root = logging.getLogger("ddanalyze")
    root.setLevel(logging.DEBUG)

    ch = logging.StreamHandler(sys.stdout)
    ch.setLevel(console_level)
    ch.setFormatter(logging.Formatter("%(asctime)s %(message)s", datefmt="%H:%M:%S"))
    root.addHandler(ch)

    fh = logging.FileHandler(log_file, encoding="utf-8")
    fh.setLevel(logging.DEBUG)
    fh.setFormatter(logging.Formatter(
        "%(asctime)s [%(levelname)-7s] %(message)s",
        datefmt="%Y-%m-%d %H:%M:%S",
    ))
    root.addHandler(fh)

    return log_file

# ─── File Utilities ───────────────────────────────────────────────────────────

def read_file(path: str) -> str:
    with open(path, "r", encoding="utf-8") as f:
        return f.read()

def write_file(path: str, content: str) -> None:
    Path(path).parent.mkdir(parents=True, exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        f.write(content)

# ─── Archive Parser ───────────────────────────────────────────────────────────

_INTERNAL_ERROR_LABELS = {
    "JSON_PARSE_ERROR", "CRITIC_JSON_PARSE_ERROR", "TIMEOUT", "FATAL_ERROR",
}

def parse_archive(archive_text: str) -> list:
    """
    Parse full_archive.txt into a list of dicts for all SUCCESS entries.
    Each entry has: iteration, date, status, analysis_type, hypothesis,
                    columns_used, code, output, evaluation.
    """
    separator = "=" * 80
    entries = []
    all_blocks = 0
    skipped_internal = 0
    skipped_non_success = 0

    for block in archive_text.split(separator):
        block = block.strip()
        if not block or "ITERATION:" not in block:
            continue

        all_blocks += 1

        status_match = re.search(r"\nSTATUS:\s*(\S+)", block)
        if not status_match:
            continue
        status = status_match.group(1).strip()
        if status.upper() in _INTERNAL_ERROR_LABELS:
            skipped_internal += 1
            continue
        if status.lower() != "success":
            skipped_non_success += 1
            continue

        entry: dict = {"status": status}

        for field, pattern in [
            ("iteration", r"ITERATION:\s*(\d+)"),
            ("date", r"DATE:\s*(.+)"),
            ("analysis_type", r"ANALYSIS TYPE:\s*(.+)"),
            ("hypothesis", r"HYPOTHESIS:\s*(.+)"),
            ("columns_used", r"COLUMNS USED:\s*(.+)"),
        ]:
            m = re.search(pattern, block)
            if m:
                entry[field] = m.group(1).strip()

        dash_sep = "-" * 80
        code_m = re.search(r"CODE:\n(.*?)\n" + re.escape(dash_sep), block, re.DOTALL)
        if code_m:
            entry["code"] = code_m.group(1).strip()

        output_m = re.search(r"OUTPUT:\n(.*?)(?:\n" + re.escape(dash_sep) + r"|\Z)", block, re.DOTALL)
        if output_m:
            entry["output"] = output_m.group(1).strip()

        eval_m = re.search(r"EVALUATION:\n(.*?)(?:\Z)", block, re.DOTALL)
        if eval_m:
            entry["evaluation"] = eval_m.group(1).strip()

        if "iteration" in entry:
            entries.append(entry)

    logger.debug(
        f"[Archive] Parsed {all_blocks} blocks: "
        f"{len(entries)} success, {skipped_non_success} failure/unknown, "
        f"{skipped_internal} internal errors"
    )
    return entries

# ─── Step A: Code Extraction ──────────────────────────────────────────────────

def extract_code_to_file(entries: list, output_path: str) -> None:
    lines = [
        "# Extracted Analysis Code\n",
        f"*Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*\n",
        f"*Total successful analyses: {len(entries)}*\n",
        "\n---\n",
    ]

    for e in entries:
        iter_num = e.get("iteration", "?")
        atype = e.get("analysis_type", "Unknown")
        hypothesis = e.get("hypothesis", "")
        columns = e.get("columns_used", "")
        code = e.get("code", "")
        output_preview = e.get("output", "")[:600]

        lines.append(f"\n## Analysis {iter_num}: {atype}\n")
        if hypothesis:
            lines.append(f"**Hypothesis:** {hypothesis}\n\n")
        if columns:
            lines.append(f"**Columns:** {columns}\n\n")
        lines.append(f"```python\n{code}\n```\n")
        if output_preview:
            lines.append(f"\n**Output (preview):**\n```\n{output_preview}\n```\n")
        lines.append("\n---\n")

    write_file(output_path, "".join(lines))
    logger.info(f"  Extracted code saved to: {output_path} ({len(entries)} analyses)")

# ─── Step B: LLM Iterative Report Generation ─────────────────────────────────

def _truncate(text: str, limit: int) -> str:
    if len(text) <= limit: return text
    return text[:limit] + "\n...[truncated]"

def list_available_graphs(graphs_folder: str) -> list:
    folder = Path(graphs_folder)
    if not folder.exists(): return []
    return [(str(p), p.name) for p in sorted(folder.glob("*.png"))]

def build_architect_prompt(entries: list, context_text: str) -> str:
    """Creates a lightweight prompt to plan the document structure."""
    summaries = []
    for e in entries:
        summaries.append(f"- ID: {e.get('iteration', '?')} | Type: {e.get('analysis_type', 'Unknown')} | Hypothesis: {e.get('hypothesis', '')}")
    
    slim_context = "\n".join(summaries)
    
    return f"""You are a senior data architect planning a final BUSINESS report for executives.
Below is the overarching context of the dataset:
{context_text}

Below is a summary of all successful data analyses performed:
{slim_context}

Your job is to organize these analyses into a logical business report. 
Create 3 to 6 logical sections (e.g., Revenue Trends, Customer Analysis, Geographic Breakdown, etc.).
Assign the relevant Analysis IDs to the section where they belong. It is okay if some IDs are left out if they aren't critical.

Respond ONLY with a valid JSON array of objects matching this exact structure (no markdown fences, no other text):
[
  {{"section_title": "Executive Summary", "iteration_ids": [], "instructions": "Write a high-level summary of the most critical findings."}},
  {{"section_title": "Revenue Dynamics", "iteration_ids": ["1", "4", "5"], "instructions": "Synthesize these analyses focusing on YoY growth and seasonality."}}
]
"""

def build_writer_system_prompt() -> str:
    return """You are a senior data analyst writing a specific section of a final BUSINESS report for executives. 
IMPORTANT: This report will become a Word document. Do NOT include Python code. Focus entirely on findings, tables, and business narrative. Do NOT use emojis.

Your job:
1. Write the section based ONLY on the provided analyses. 
2. Write a clear, business-readable title (no jargon).
3. MANDATORY: Include a year-by-year comparison table whenever the finding spans multiple periods.
   Always show 3–4 years + LTM (Last-Twelve-Months). Include a CAGR row/column when you have 3+ years of data.
   
   CRITICAL TABLE FORMATTING: 
   Time periods must ALWAYS be columns (left-to-right). Metrics must ALWAYS be rows (top-to-bottom).
   Format tables strictly like this:
       | Metric      | FY2021 | FY2022 | FY2023 | LTM    | CAGR   |
       |-------------|--------|--------|--------|--------|--------|
       | Revenue     | $X     | $X     | $X     | $X     | XX%    |
       | Customers   | X      | X      | X      | X      | XX%    |
       | YoY %       | —      | +X%    | +X%    | +X%    | —      |

4. If a graph is relevant to the data, reference it with EXACTLY this syntax: [GRAPH: filename.png]
5. TONE: Formal, senior consulting register. Avoid conversational phrases. "This pattern is consistent with..."
6. Define any internal product acronyms the first time you use them.

Output ONLY the markdown for this specific section. Do not include an introductory or concluding remark outside of the requested report text."""

def build_writer_user_prompt(section_title: str, instructions: str, relevant_entries: list, available_graphs: list) -> str:
    parts = []
    for e in relevant_entries:
        parts.append(f"""--- ANALYSIS {e.get('iteration', '?')} ---
Type: {e.get('analysis_type', '')}
Hypothesis: {e.get('hypothesis', '')}
Output:\n{_truncate(e.get('output', ''), 2000)}
Evaluation:\n{_truncate(e.get('evaluation', ''), 800)}""")
    
    analyses_text = "\n\n".join(parts) if parts else "No specific data provided for this section. Use overarching knowledge."
    
    graphs_section = ""
    if available_graphs:
        graphs_section = "Available Graphs:\n" + "\n".join([f"- {name}" for _, name in available_graphs])

    return f"""SECTION TO WRITE: {section_title}
INSTRUCTIONS: {instructions}

{graphs_section}

DATA FOR THIS SECTION:
{analyses_text}

Please generate the markdown for this section now."""

# ─── Step C: Word Document Builder ───────────────────────────────────────────

def _parse_table_row(line: str) -> list:
    cells = [c.strip() for c in line.split("|")]
    return [c for c in cells if c] 

def _is_separator_row(line: str) -> bool:
    return bool(re.match(r"^\|[\s\-:|]+\|$", line.strip()))

def _set_cell_shading(cell, hex_color: str) -> None:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)

def _add_markdown_table_to_doc(doc, header_line: str, data_lines: list) -> None:
    from docx.shared import Pt, RGBColor
    HEADER_BG  = "046A38"
    HEADER_FG  = RGBColor(0xFF, 0xFF, 0xFF)
    BODY_FONT  = "Arial"

    header_cells = _parse_table_row(header_line)
    if not header_cells: return
    n_cols = len(header_cells)
    
    data_rows = []
    for line in data_lines:
        row = _parse_table_row(line)
        if row:
            while len(row) < n_cols: row.append("")
            data_rows.append(row[:n_cols])

    all_rows = [header_cells] + data_rows
    if not all_rows: return

    table = doc.add_table(rows=len(all_rows), cols=n_cols)
    try: table.style = "Table Grid"
    except Exception: pass

    for row_idx, row_data in enumerate(all_rows):
        for col_idx, cell_text in enumerate(row_data):
            if col_idx >= n_cols: continue
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = "" 
            para = cell.paragraphs[0]
            run = para.add_run(cell_text)
            run.font.name = BODY_FONT
            run.font.size = Pt(10)

            if row_idx == 0:
                _set_cell_shading(cell, HEADER_BG)
                run.bold = True
                run.font.color.rgb = HEADER_FG
    doc.add_paragraph()

def _add_formatted_paragraph(doc, text: str):
    if "**" not in text:
        doc.add_paragraph(text)
        return
    para = doc.add_paragraph()
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            para.add_run(part)

def _add_deloitte_heading(doc, text: str, level: int, title_rgb, font: str) -> None:
    para = doc.add_heading(text, level=level)
    for run in para.runs:
        run.font.color.rgb = title_rgb
        run.font.name = font

def build_word_document(report_md: str, graphs_folder: str, output_path: str) -> None:
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
    except ImportError:
        logger.warning("\n[WARNING] python-docx is not installed. Markdown saved only.\n")
        return

    t0 = time.time()
    doc = Document()
    graphs_embedded, graphs_missing, tables_added = 0, 0, 0

    for section in doc.sections:
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)

    _DELOITTE_TITLE_RGB = RGBColor(0x26, 0x89, 0x0D)
    _ARIAL = "Arial"
    try:
        normal_style = doc.styles["Normal"]
        normal_style.font.name = _ARIAL
        normal_style.font.size = Pt(10)
    except Exception: pass

    lines = report_md.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("### "):
            _add_deloitte_heading(doc, stripped[4:], 3, _DELOITTE_TITLE_RGB, _ARIAL)
        elif stripped.startswith("## "):
            _add_deloitte_heading(doc, stripped[3:], 2, _DELOITTE_TITLE_RGB, _ARIAL)
        elif stripped.startswith("# "):
            _add_deloitte_heading(doc, stripped[2:], 1, _DELOITTE_TITLE_RGB, _ARIAL)
        elif stripped.startswith("|") and i + 1 < len(lines) and _is_separator_row(lines[i + 1]):
            header_line = stripped
            i += 2 
            data_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                data_lines.append(lines[i].strip())
                i += 1
            _add_markdown_table_to_doc(doc, header_line, data_lines)
            tables_added += 1
            continue 
        elif "[GRAPH:" in stripped:
            match = re.search(r"\[GRAPH:\s*([\w\-\.]+\.png)\]", stripped, re.IGNORECASE)
            if match:
                filename = match.group(1)
                img_path = Path(graphs_folder) / filename
                if img_path.exists():
                    try:
                        doc.add_picture(str(img_path), width=Inches(6.0))
                        caption = re.sub(r"\[GRAPH:.*?\]", "", stripped, flags=re.IGNORECASE).strip()
                        if caption:
                            p = doc.add_paragraph(caption)
                            for run in p.runs: run.italic = True
                        graphs_embedded += 1
                    except Exception as e:
                        logger.warning(f"[DocBuilder] Failed to embed {filename}: {e}")
                else:
                    doc.add_paragraph(f"[Graph not found: {filename}]")
                    graphs_missing += 1
            else:
                _add_formatted_paragraph(doc, stripped)
        elif stripped.startswith("!["):
            match = re.search(r"!\[([^\]]*)\]\(([^\)]+)\)", stripped)
            if match:
                img_path = match.group(2)
                if Path(img_path).exists():
                    try:
                        doc.add_picture(img_path, width=Inches(6.0))
                        graphs_embedded += 1
                    except Exception: pass
                else:
                    doc.add_paragraph(f"[Image not found: {img_path}]")
            else:
                _add_formatted_paragraph(doc, stripped)
        elif stripped.startswith("---") and stripped == "-" * len(stripped):
            doc.add_paragraph() 
        elif stripped:
            _add_formatted_paragraph(doc, stripped)
        i += 1

    doc.save(output_path)
    logger.info(f"  Word document saved to: {output_path} (Built in {time.time() - t0:.1f}s)")

# ─── Main Orchestrator ────────────────────────────────────────────────────────

def main() -> None:
    # Safely load config inside main to prevent module-level crashes
    try:
        config = yaml.safe_load(read_file("config.yaml"))
    except FileNotFoundError:
        logger.warning("config.yaml not found! Using default settings.")
        config = {}

    # Ensure a valid Anthropic model is used
    model = config.get("model", "claude-4-6") 
    max_tokens = config.get("phase2_max_tokens", 20000)
    
    archive_path = config.get("archive_file", "full_archive.txt")
    context_path = config.get("active_context_file", "active_context.md")
    graphs_folder = config.get("graphs_folder", "workspace/graphs")
    
    extracted_code_path = "extracted_code.md"
    report_md_path = "final_report.md"
    report_docx_path = "final_report.docx"

    log_file = setup_logging(debug=config.get("debug_logging", False))
    logger.info("=" * 60)
    logger.info("Phase 2 — Final Report Generator (Iterative Mode)")
    logger.info("=" * 60)

    if not Path(archive_path).exists():
        logger.error(f"ERROR: {archive_path} not found. Exiting.")
        sys.exit(1)

    archive_text = read_file(archive_path)
    context_text = read_file(context_path) if Path(context_path).exists() else ""
    entries = parse_archive(archive_text)
    
    if not entries:
        logger.error("No successful analyses to report on. Exiting.")
        sys.exit(0)

    # STEP A: Extract Code
    logger.info("\n[Step A] Extracting code...")
    extract_code_to_file(entries, extracted_code_path)

    # STEP B: Iterative LLM Generation
    logger.info(f"\n[Step B] Initiating 'Outline & Draft' architecture...")
    available_graphs = list_available_graphs(graphs_folder)
    client = anthropic.Anthropic(api_key=API_KEY)

    # B.1: The Architect (Plan the report)
    logger.info("  -> Calling Architect for Report Outline...")
    architect_prompt = build_architect_prompt(entries, context_text)
    
    architect_response = client.messages.create(
        model=model,
        max_tokens=2048,
        messages=[{"role": "user", "content": architect_prompt}]
    )
    
    # Parse JSON strictly, removing markdown fences if the LLM added them
    raw_json = architect_response.content[0].text.strip()
    raw_json = re.sub(r"^```json", "", raw_json)
    raw_json = re.sub(r"```$", "", raw_json).strip()
    
    try:
        outline = json.loads(raw_json)
        logger.info(f"  -> Architect mapped {len(outline)} sections.")
    except json.JSONDecodeError as e:
        logger.error(f"Failed to parse Architect JSON: {e}\nRaw Output:\n{raw_json}")
        sys.exit(1)

    # B.2: The Writer (Iterative Drafting)
    full_markdown_report = "# Final Business Report\n\n"
    system_prompt = build_writer_system_prompt()

    for idx, section in enumerate(outline, 1):
        title = section.get("section_title", f"Section {idx}")
        target_ids = section.get("iteration_ids", [])
        instructions = section.get("instructions", "")
        
        # Filter entries down to only what is needed for this specific section
        relevant_entries = [e for e in entries if e.get("iteration") in target_ids]
        
        logger.info(f"  -> Drafting Section {idx}/{len(outline)}: '{title}' (using {len(relevant_entries)} analyses)...")
        
        user_prompt = build_writer_user_prompt(title, instructions, relevant_entries, available_graphs)
        
        writer_response = client.messages.create(
            model=model,
            max_tokens=max_tokens,
            system=system_prompt,
            messages=[{"role": "user", "content": user_prompt}]
        )
        
        section_content = writer_response.content[0].text.strip()
        full_markdown_report += f"## {title}\n\n{section_content}\n\n---\n\n"

    write_file(report_md_path, full_markdown_report)
    logger.info(f"  -> Iterative drafting complete. Markdown saved to: {report_md_path}")

    # STEP C: Build Word Doc
    logger.info(f"\n[Step C] Building Word document ({report_docx_path})...")
    build_word_document(full_markdown_report, graphs_folder, report_docx_path)

    logger.info("\n" + "=" * 60)
    logger.info("Phase 2 complete. All files saved successfully.")

if __name__ == "__main__":
    main()