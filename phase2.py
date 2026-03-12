"""
Phase 2 — Final Report Generator

Three-step process:
  Step A: Extract all successful analysis code → extracted_code.md
  Step B: LLM synthesises findings into a structured markdown report → final_report.md
  Step C: Convert the markdown report to a Word document → final_report.docx
          (embeds tables as Word tables and inserts graph images)

Run this manually after loop.py finishes.
"""

import logging
import os
import re
import sys
import time
from datetime import datetime
from pathlib import Path
from typing import Optional

import anthropic
import yaml

# Change to the directory containing this script so relative paths work
os.chdir(Path(__file__).parent)

from dotenv import load_dotenv

load_dotenv()
API_KEY = os.getenv('API_KEY')

# ─── Logging Setup ────────────────────────────────────────────────────────────

logger = logging.getLogger("ddanalyze.phase2")

def setup_logging(debug: bool = False, log_dir: str = "logs") -> Path:
    """Configure console + file logging. Returns the path of the log file."""
    Path(log_dir).mkdir(exist_ok=True)
    log_file = Path(log_dir) / f"phase2_{datetime.now().strftime('%Y%m%d_%H%M%S')}.log"

    console_level = logging.DEBUG if debug else logging.INFO

    root = logging.getLogger("ddanalyze")
    root.setLevel(logging.DEBUG)

    ch = logging.StreamHandler(sys.stdout)
    ch.setLevel(console_level)
    ch.setFormatter(logging.Formatter("%(asctime)s %(message)s", datefmt="%H:%M:%S"))
    root.addHandler(ch)

    fh = logging.FileHandler(log_file, encoding="utf-8")
    fh.setLevel(logging.DEBUG)
    fh.setFormatter(logging.Formatter(
        "%(asctime)s [%(levelname)-7s] %(message)s",
        datefmt="%Y-%m-%d %H:%M:%S",
    ))
    root.addHandler(fh)

    return log_file

# ─── System Prompt ────────────────────────────────────────────────────────────

PHASE2_SYSTEM_PROMPT = """You are a senior data analyst preparing a final BUSINESS report for executives. \
You have access to a complete log of autonomous data analyses performed on a company dataset.

IMPORTANT: This report will become a Word document for business readers. Do NOT include Python code.
Code has already been extracted separately. Focus entirely on findings, tables, and business narrative.

Your job:
1. Select the most valuable findings — those that best explain the business, its customers,
   its revenue dynamics, and its key patterns.

2. For each selected finding, write:
   - A clear, business-readable title (no jargon)
   - MANDATORY: A year-by-year comparison table whenever the finding spans multiple periods.
     Always show 3–4 years + LTM (Last-Twelve-Months) if the latest year is incomplete.
     Always include a CAGR row when you have 3+ years of data.
     Format tables as proper markdown tables:
       | Year   | Revenue  | Customers | YoY %  |
       |--------|----------|-----------|--------|
       | FY2021 | $X       | X         | —      |
       | FY2022 | $X       | X         | +X%    |
       | FY2023 | $X       | X         | +X%    |
       | LTM    | $X       | X         | +X%    |
       | CAGR   |          |           | XX%    |
   - A 3–5 sentence explanation of what was found and why it matters for the business.
   - If a graph was generated for this analysis, reference it with EXACTLY this syntax:
       [GRAPH: filename.png]
     (use the filename from the Available Graphs list)

3. Organize findings into logical sections (e.g. Revenue Trends, Customer Analysis,
   Concentration & Risk, Seasonality & Timing, Geographic / Segment Breakdown).

4. End with a one-page Executive Summary: the 5–7 most important things to know about
   this business from the data. Use clear bullet points.

5. TONE — Write in a formal, senior consulting register throughout. Avoid conversational,
   didactic, or explanatory phrases (e.g. "This is exactly how wholesale ordering works:",
   "This makes sense because", "It's worth noting that"). Instead, use professional
   formulations such as "This pattern is consistent with...", "This trend reflects...",
   "This dynamic is in line with...", "As expected in the X segment, ...".
   Never explain obvious industry mechanics as if educating the reader.

6. TERM INTRODUCTION — When a specific product model name, customer segment label, or
   internal category name is referenced for the first time in a section, provide brief
   contextual identification so that a reader unfamiliar with the company's internal
   naming conventions can follow. Embed this naturally in the prose — for example:
   "True Rising — the brand's highest-revenue running model — posted..." or
   "the TRIBE line (lifestyle and fashion-oriented footwear)". Do not assume the reader
   knows internal product names or segment definitions.

Write for a senior business audience. Avoid statistical jargon. Assume readers are familiar
with industry concepts but not with this company's internal product names or segment labels.
Output as clean markdown. Do NOT include Python code blocks."""

# ─── Iterative Pipeline Prompts ──────────────────────────────────────────────

ARCHITECT_SYSTEM_PROMPT = """\
You are a report architect. You will receive the full knowledge base and analysis log from an \
autonomous data-analysis pipeline. Your ONLY job is to design the structure of the final report.

Output a single valid JSON object — no markdown fences, no commentary, no prose.

The JSON schema:
{
  "report_title": "string — title for the overall report",
  "sections": [
    {
      "section_number": int,
      "title": "string — business-readable section heading",
      "description": "string — 1-2 sentence scope of this section",
      "iterations": [int, ...],   // which analysis iteration numbers belong here
      "graphs": ["filename.png", ...],  // which graphs to embed in this section
      "guidance": "string — writing instructions for the section author"
    }
  ],
  "executive_summary_guidance": "string — key points the exec summary must cover"
}

Rules:
1. Every successful analysis iteration must appear in EXACTLY ONE section.
2. Order sections by business logic: start with high-level revenue/growth, move to product, \
then customer/channel, then risk/concentration, then any remaining themes.
3. Aim for 4-8 sections. Merge small related analyses; split large heterogeneous ones.
4. In "guidance", specify what tables to include (year-by-year, top-N, etc.), what narrative \
angle to take, and what comparisons to draw.
5. Map graphs to sections based on their filenames and the analyses that generated them.
6. The executive_summary_guidance should list the 5-7 most important business takeaways \
the summary must cover.
7. Output ONLY the JSON object. No extra text before or after."""

SECTION_WRITER_SYSTEM_PROMPT = """\
You are a senior data analyst writing ONE section of a business report for executives.

You will receive:
- The overall knowledge base (for grounding and cross-references)
- The specific analyses assigned to YOUR section (2-5 analyses)
- Writing guidance from the report architect
- The list of graphs available for your section

Your job for THIS SECTION ONLY:
1. Write a clear ## section heading followed by the detailed narrative.
2. For each key finding, include:
   - A business-readable sub-heading (### level)
   - MANDATORY year-by-year comparison tables when findings span multiple periods.
     Always show 3-4 years + LTM if the latest year is incomplete.
     Always include a CAGR row when you have 3+ years of data.
     Format as proper markdown tables.
   - A 3-5 sentence explanation of what was found and why it matters.
   - Graph references using EXACTLY: [GRAPH: filename.png]
3. Be thorough and detailed — you are writing ONLY this section, so give it full attention.
4. Do NOT write an executive summary — that will be handled separately.
5. Do NOT include Python code.

TONE — Write in a formal, senior consulting register. Avoid conversational or didactic phrases. \
Use professional formulations: "This pattern is consistent with...", "This trend reflects...", etc.

TERM INTRODUCTION — When a specific product model name, customer segment label, or internal \
category name is referenced for the first time, provide brief contextual identification so that \
a reader unfamiliar with the company can follow. Embed this naturally in the prose.

Output clean markdown for this section only. Start with ## heading."""

EXEC_SUMMARY_SYSTEM_PROMPT = """\
You are a senior data analyst writing the Executive Summary for a business due-diligence report.

You will receive:
- The section titles and key data tables from the completed report
- Guidance on what the summary must cover
- The overall knowledge base

Write a one-page Executive Summary:
1. Start with ## Executive Summary
2. Open with a 2-3 sentence overview of the business and dataset scope.
3. Then list 5-7 bullet points — the most important things to know about this business.
   Each bullet should be specific (include numbers) and actionable.
4. Close with a brief forward-looking paragraph on key risks and opportunities.

TONE — Formal, senior consulting register. No jargon. Specific numbers over vague statements.
Output clean markdown. Do NOT include Python code."""

# ─── File Utilities ───────────────────────────────────────────────────────────

def read_file(path: str) -> str:
    with open(path, "r", encoding="utf-8") as f:
        return f.read()

def write_file(path: str, content: str) -> None:
    Path(path).parent.mkdir(parents=True, exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        f.write(content)

# ─── Archive Parser ───────────────────────────────────────────────────────────

_INTERNAL_ERROR_LABELS = {
    "JSON_PARSE_ERROR", "CRITIC_JSON_PARSE_ERROR", "TIMEOUT", "FATAL_ERROR",
}
config = yaml.safe_load(read_file("config.yaml"))
PHASE2_MAX_TOKENS = config.get("analyst_max_tokens", 16384)


def parse_archive(archive_text: str) -> list:
    """
    Parse full_archive.txt into a list of dicts for all SUCCESS entries.
    Each entry has: iteration, date, status, analysis_type, hypothesis,
                    columns_used, code, output, evaluation.
    """
    separator = "=" * 80
    entries = []
    all_blocks = 0
    skipped_internal = 0
    skipped_non_success = 0

    for block in archive_text.split(separator):
        block = block.strip()
        if not block or "ITERATION:" not in block:
            continue

        all_blocks += 1

        # Skip internal error blocks
        status_match = re.search(r"\nSTATUS:\s*(\S+)", block)
        if not status_match:
            continue
        status = status_match.group(1).strip()
        if status.upper() in _INTERNAL_ERROR_LABELS:
            skipped_internal += 1
            continue
        if status.lower() != "success":
            skipped_non_success += 1
            continue

        entry: dict = {"status": status}

        # Simple line-level fields
        for field, pattern in [
            ("iteration", r"ITERATION:\s*(\d+)"),
            ("date", r"DATE:\s*(.+)"),
            ("analysis_type", r"ANALYSIS TYPE:\s*(.+)"),
            ("hypothesis", r"HYPOTHESIS:\s*(.+)"),
            ("columns_used", r"COLUMNS USED:\s*(.+)"),
        ]:
            m = re.search(pattern, block)
            if m:
                entry[field] = m.group(1).strip()

        # Code block (between "CODE:\n" and the next "---" separator)
        dash_sep = "-" * 80
        code_m = re.search(
            r"CODE:\n(.*?)\n" + re.escape(dash_sep), block, re.DOTALL
        )
        if code_m:
            entry["code"] = code_m.group(1).strip()

        # Output block (between "OUTPUT:\n" and next "---" separator or end)
        output_m = re.search(
            r"OUTPUT:\n(.*?)(?:\n" + re.escape(dash_sep) + r"|\Z)", block, re.DOTALL
        )
        if output_m:
            entry["output"] = output_m.group(1).strip()

        # Evaluation block (from "EVALUATION:\n" to end)
        eval_m = re.search(r"EVALUATION:\n(.*?)(?:\Z)", block, re.DOTALL)
        if eval_m:
            entry["evaluation"] = eval_m.group(1).strip()

        if "iteration" in entry:
            entries.append(entry)

    logger.debug(
        f"[Archive] Parsed {all_blocks} blocks: "
        f"{len(entries)} success, {skipped_non_success} failure/unknown, "
        f"{skipped_internal} internal errors"
    )
    return entries

# ─── Step A: Code Extraction ──────────────────────────────────────────────────

def extract_code_to_file(entries: list, output_path: str) -> None:
    """
    Generate extracted_code.md with all successful analysis code blocks,
    organised by analysis type and iteration number.
    """
    lines = [
        "# Extracted Analysis Code\n",
        f"*Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*\n",
        f"*Total successful analyses: {len(entries)}*\n",
        "\n---\n",
    ]

    for e in entries:
        iter_num = e.get("iteration", "?")
        atype = e.get("analysis_type", "Unknown")
        hypothesis = e.get("hypothesis", "")
        columns = e.get("columns_used", "")
        code = e.get("code", "")
        output_preview = e.get("output", "")[:600]

        lines.append(f"\n## Analysis {iter_num}: {atype}\n")
        if hypothesis:
            lines.append(f"**Hypothesis:** {hypothesis}\n\n")
        if columns:
            lines.append(f"**Columns:** {columns}\n\n")
        lines.append(f"```python\n{code}\n```\n")
        if output_preview:
            lines.append(f"\n**Output (preview):**\n```\n{output_preview}\n```\n")
        lines.append("\n---\n")

    write_file(output_path, "".join(lines))
    logger.info(f"  Extracted code saved to: {output_path} ({len(entries)} analyses)")

# ─── Step B: LLM Report Generation ───────────────────────────────────────────

def _truncate(text: str, limit: int) -> str:
    if len(text) <= limit:
        return text
    return text[:limit] + "\n...[truncated]"

def list_available_graphs(graphs_folder: str) -> list:
    """Return list of (filepath, filename) for all PNG files in the graphs folder."""
    folder = Path(graphs_folder)
    if not folder.exists():
        return []
    return [(str(p), p.name) for p in sorted(folder.glob("*.png"))]

def build_report_prompt(entries: list, context_text: str, available_graphs: list) -> str:
    parts = []
    for e in entries:
        part = f"""--- ANALYSIS {e.get('iteration', '?')} ---
Type       : {e.get('analysis_type', '')}
Hypothesis : {e.get('hypothesis', '')}
Columns    : {e.get('columns_used', '')}

Output:
{_truncate(e.get('output', ''), 2000)}

Evaluation:
{_truncate(e.get('evaluation', ''), 800)}
"""
        parts.append(part)

    analyses_text = "\n\n".join(parts)

    graphs_section = ""
    if available_graphs:
        graphs_section = "\n\n## Available Graphs\n"
        graphs_section += "The following graph files were saved during the analysis loop.\n"
        graphs_section += "Reference them in the report using [GRAPH: filename.png] syntax.\n\n"
        for filepath, filename in available_graphs:
            graphs_section += f"- {filename}\n"

    return f"""## Final State of the Knowledge Base

{context_text}
{graphs_section}
---

## Complete Log of Successful Analyses ({len(entries)} total)

{analyses_text}

---

Please generate a comprehensive final business report based on all of the above.
Select the most valuable analyses, group them into logical sections, include year-by-year
tables wherever relevant, embed graph references using [GRAPH: filename.png] syntax,
and conclude with a concise executive summary. Do NOT include Python code."""


# ─── Step B (Iterative): Architect → Filter → Writer → Assembler ─────────────

import json


def build_architect_prompt(entries: list, context_text: str, available_graphs: list) -> str:
    """Build the user prompt for B.1 — the Architect generates a JSON outline."""
    analyses_summary = []
    for e in entries:
        analyses_summary.append(
            f"- Iteration {e.get('iteration', '?')}: "
            f"[{e.get('analysis_type', 'Unknown')}] {e.get('hypothesis', '')}"
            f"\n  Output excerpt: {_truncate(e.get('output', ''), 800)}"
            f"\n  Evaluation excerpt: {_truncate(e.get('evaluation', ''), 400)}"
        )
    analyses_text = "\n".join(analyses_summary)

    graphs_text = ""
    if available_graphs:
        graphs_text = "\n## Available Graphs\n" + "\n".join(
            f"- {fname}" for _, fname in available_graphs
        )

    return f"""## Knowledge Base
{context_text}
{graphs_text}

## Successful Analyses ({len(entries)} total)
{analyses_text}

---

Design the report structure. Output a single JSON object following the schema in your instructions.
Every iteration number must appear in exactly one section. Map graphs to the sections where they are most relevant."""


def parse_outline_json(raw_text: str) -> dict:
    """Parse the Architect's JSON output, handling markdown fences if present."""
    text = raw_text.strip()
    # Strip markdown code fences if the model wrapped its output
    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?\s*\n?", "", text)
        text = re.sub(r"\n?```\s*$", "", text)

    outline = json.loads(text)

    # Basic validation
    if "sections" not in outline or not isinstance(outline["sections"], list):
        raise ValueError("Outline JSON missing 'sections' array")
    for sec in outline["sections"]:
        if "iterations" not in sec or "title" not in sec:
            raise ValueError(f"Section missing required fields: {sec}")
        sec.setdefault("section_number", 0)
        sec.setdefault("description", "")
        sec.setdefault("graphs", [])
        sec.setdefault("guidance", "")
    outline.setdefault("report_title", "Final Analysis Report")
    outline.setdefault("executive_summary_guidance", "")

    return outline


def filter_entries_for_section(entries: list, iteration_ids: list) -> list:
    """B.2 — Return only entries whose iteration number is in the given list."""
    id_set = {str(i) for i in iteration_ids}
    return [e for e in entries if str(e.get("iteration", "")) in id_set]


def build_section_prompt(
    section: dict,
    section_entries: list,
    context_text: str,
    section_graphs: list,
    previous_titles: list,
) -> str:
    """Build the user prompt for B.3 — writing one section."""
    # Detailed analysis data for just this section's entries
    parts = []
    for e in section_entries:
        part = f"""--- ANALYSIS {e.get('iteration', '?')} ---
Type       : {e.get('analysis_type', '')}
Hypothesis : {e.get('hypothesis', '')}
Columns    : {e.get('columns_used', '')}

Output:
{_truncate(e.get('output', ''), 3000)}

Evaluation:
{_truncate(e.get('evaluation', ''), 1200)}
"""
        parts.append(part)
    analyses_text = "\n\n".join(parts)

    graphs_text = ""
    if section_graphs:
        graphs_text = "\n## Graphs Available for This Section\n" + "\n".join(
            f"- {fname}" for _, fname in section_graphs
        )

    prev_text = ""
    if previous_titles:
        prev_text = (
            "\n## Sections Already Written (avoid repeating their content)\n"
            + "\n".join(f"- {t}" for t in previous_titles)
        )

    return f"""## Section to Write
Title: {section['title']}
Description: {section.get('description', '')}
Architect Guidance: {section.get('guidance', '')}
{prev_text}

## Knowledge Base (for context and grounding)
{context_text}
{graphs_text}

## Analyses Assigned to This Section ({len(section_entries)} total)

{analyses_text}

---

Write this section now. Start with ## {section['title']} and be thorough and detailed.
Include year-by-year tables wherever the data supports it. Reference graphs using [GRAPH: filename.png]."""


def build_executive_summary_prompt(
    outline: dict,
    section_markdowns: list,
    context_text: str,
) -> str:
    """Build the user prompt for the final exec summary call."""
    # Extract section titles and any tables from the written sections
    section_digest = []
    for i, md in enumerate(section_markdowns):
        # Grab heading and any tables
        lines = md.split("\n")
        digest_lines = []
        for line in lines:
            stripped = line.strip()
            if stripped.startswith("#") or stripped.startswith("|"):
                digest_lines.append(stripped)
        section_digest.append("\n".join(digest_lines) if digest_lines else f"Section {i+1}")

    digest_text = "\n\n".join(section_digest)

    return f"""## Architect's Summary Guidance
{outline.get('executive_summary_guidance', 'Summarize the 5-7 most important findings.')}

## Section Headings and Key Tables from the Completed Report
{digest_text}

## Knowledge Base
{context_text}

---

Write the Executive Summary now. Start with ## Executive Summary."""


def _llm_call(
    client,
    model: str,
    system_prompt: str,
    user_message: str,
    max_tokens: int,
    label: str = "",
) -> tuple:
    """Make a single streaming LLM call. Returns (text, usage)."""
    logger.info(f"  [{label}] Sending request ({len(user_message):,}ch prompt, max_tokens={max_tokens})")

    text = ""
    with client.messages.stream(
        model=model,
        max_tokens=max_tokens,
        system=system_prompt,
        messages=[{"role": "user", "content": user_message}],
    ) as stream:
        for chunk in stream.text_stream:
            print(chunk, end="", flush=True)
        response = stream.get_final_message()
    print()  # newline after streaming

    for block in response.content:
        if block.type == "text":
            text = block.text
            break

    usage = response.usage
    logger.info(
        f"  [{label}] Done — {len(text):,}ch | "
        f"tokens in={usage.input_tokens:,} out={usage.output_tokens:,}"
    )
    return text, usage


def generate_report_iterative(
    client,
    model: str,
    entries: list,
    context_text: str,
    available_graphs: list,
    config: dict,
) -> str:
    """
    Iterative report generation pipeline: Architect → Filter → Writer → Assembler.
    Returns the final assembled markdown report.
    """
    architect_max_tokens = config.get("phase2_architect_max_tokens", 4000)
    section_max_tokens = config.get("phase2_section_max_tokens", 8000)
    summary_max_tokens = config.get("phase2_summary_max_tokens", 4000)

    total_usage = {"input": 0, "output": 0}

    # ── B.1: The Architect ────────────────────────────────────────────────────
    logger.info("\n[B.1] The Architect — generating report outline...")
    architect_prompt = build_architect_prompt(entries, context_text, available_graphs)

    architect_raw, usage = _llm_call(
        client, model, ARCHITECT_SYSTEM_PROMPT, architect_prompt,
        architect_max_tokens, label="B.1 Architect",
    )
    total_usage["input"] += usage.input_tokens
    total_usage["output"] += usage.output_tokens

    # Parse the JSON outline
    try:
        outline = parse_outline_json(architect_raw)
    except (json.JSONDecodeError, ValueError) as e:
        logger.warning(f"[B.1] JSON parse failed ({e}), retrying with reinforcement...")
        retry_prompt = (
            architect_prompt
            + "\n\nIMPORTANT: Your previous response was not valid JSON. "
            "Output ONLY a valid JSON object, no markdown fences, no commentary."
        )
        architect_raw, usage = _llm_call(
            client, model, ARCHITECT_SYSTEM_PROMPT, retry_prompt,
            architect_max_tokens, label="B.1 Architect (retry)",
        )
        total_usage["input"] += usage.input_tokens
        total_usage["output"] += usage.output_tokens
        try:
            outline = parse_outline_json(architect_raw)
        except (json.JSONDecodeError, ValueError) as e2:
            logger.error(f"[B.1] Retry also failed ({e2}). Falling back to legacy single-call.")
            return None  # Signal caller to use legacy path

    logger.info(
        f"[B.1] Outline: {len(outline['sections'])} sections, "
        f"title=\"{outline.get('report_title', 'N/A')}\""
    )
    for sec in outline["sections"]:
        logger.info(
            f"       Section {sec['section_number']}: {sec['title']} "
            f"(iterations={sec['iterations']}, graphs={sec.get('graphs', [])})"
        )

    # ── B.2 + B.3: Filter & Write each section ───────────────────────────────
    section_markdowns = []
    previous_titles = []

    for sec_idx, section in enumerate(outline["sections"]):
        sec_num = sec_idx + 1
        logger.info(f"\n[B.3] Writing section {sec_num}/{len(outline['sections'])}: {section['title']}")

        # B.2: Filter entries for this section
        section_entries = filter_entries_for_section(entries, section["iterations"])
        section_graphs = [
            g for g in available_graphs if g[1] in section.get("graphs", [])
        ]

        if not section_entries:
            logger.warning(f"  [B.3] No entries found for section {sec_num}, skipping.")
            section_markdowns.append(f"## {section['title']}\n\n*No analyses available for this section.*\n")
            previous_titles.append(section["title"])
            continue

        logger.info(
            f"  [B.2] Filtered: {len(section_entries)} analyses, "
            f"{len(section_graphs)} graphs"
        )

        # B.3: Write the section
        section_prompt = build_section_prompt(
            section, section_entries, context_text, section_graphs, previous_titles,
        )

        section_md, usage = _llm_call(
            client, model, SECTION_WRITER_SYSTEM_PROMPT, section_prompt,
            section_max_tokens, label=f"B.3 Section {sec_num}",
        )
        total_usage["input"] += usage.input_tokens
        total_usage["output"] += usage.output_tokens

        section_markdowns.append(section_md)
        previous_titles.append(section["title"])

    # ── B.3 (final): Executive Summary ────────────────────────────────────────
    logger.info(f"\n[B.3] Writing Executive Summary...")
    summary_prompt = build_executive_summary_prompt(outline, section_markdowns, context_text)

    summary_md, usage = _llm_call(
        client, model, EXEC_SUMMARY_SYSTEM_PROMPT, summary_prompt,
        summary_max_tokens, label="B.3 Exec Summary",
    )
    total_usage["input"] += usage.input_tokens
    total_usage["output"] += usage.output_tokens

    # ── B.4: Assemble ─────────────────────────────────────────────────────────
    logger.info(f"\n[B.4] Assembling final report...")
    report_title = outline.get("report_title", "Final Analysis Report")
    final_parts = [f"# {report_title}\n"]
    for section_md in section_markdowns:
        final_parts.append(section_md)
    final_parts.append(summary_md)

    final_report = "\n\n---\n\n".join(final_parts)

    logger.info(
        f"[B.4] Assembly complete — {len(final_report):,}ch | "
        f"Total tokens: in={total_usage['input']:,} out={total_usage['output']:,} "
        f"combined={total_usage['input'] + total_usage['output']:,}"
    )

    return final_report


# ─── Step C: Word Document Builder ───────────────────────────────────────────

def _parse_table_row(line: str) -> list:
    """Parse a markdown table row string into a list of cell strings."""
    cells = [c.strip() for c in line.split("|")]
    return [c for c in cells if c]  # Remove empty strings from leading/trailing |

def _is_separator_row(line: str) -> bool:
    """Return True if the line is a markdown table separator (|---|---|)."""
    return bool(re.match(r"^\|[\s\-:|]+\|$", line.strip()))

def _set_cell_shading(cell, hex_color: str) -> None:
    """Apply a solid background fill to a table cell using direct XML manipulation."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)


def _add_markdown_table_to_doc(doc, header_line: str, data_lines: list) -> None:
    """Convert parsed markdown table lines into a python-docx Word table.

    Header row: Deloitte dark-green (#046A38) background, white bold Arial text.
    Body rows: Arial text, normal weight.
    """
    from docx.shared import Pt, RGBColor

    # Deloitte table header colours
    HEADER_BG  = "046A38"   # dark green (no leading #)
    HEADER_FG  = RGBColor(0xFF, 0xFF, 0xFF)   # white
    BODY_FONT  = "Arial"

    header_cells = _parse_table_row(header_line)
    if not header_cells:
        return

    n_cols = len(header_cells)
    data_rows = []
    for line in data_lines:
        row = _parse_table_row(line)
        if row:
            while len(row) < n_cols:
                row.append("")
            data_rows.append(row[:n_cols])

    all_rows = [header_cells] + data_rows
    if not all_rows:
        return

    table = doc.add_table(rows=len(all_rows), cols=n_cols)
    try:
        table.style = "Table Grid"
    except Exception:
        pass

    for row_idx, row_data in enumerate(all_rows):
        for col_idx, cell_text in enumerate(row_data):
            if col_idx >= n_cols:
                continue
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = ""  # clear before adding a styled run

            para = cell.paragraphs[0]
            run = para.add_run(cell_text)
            run.font.name = BODY_FONT
            run.font.size = Pt(10)

            if row_idx == 0:
                # Header: dark-green background, white bold text
                _set_cell_shading(cell, HEADER_BG)
                run.bold = True
                run.font.color.rgb = HEADER_FG

    doc.add_paragraph()

def _add_formatted_paragraph(doc, text: str):
    """Add a paragraph with basic **bold** support."""
    if "**" not in text:
        doc.add_paragraph(text)
        return

    para = doc.add_paragraph()
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            para.add_run(part)

def _add_deloitte_heading(doc, text: str, level: int, title_rgb, font: str) -> None:
    """Add a Word heading styled with the Deloitte green color and Arial font."""
    para = doc.add_heading(text, level=level)
    for run in para.runs:
        run.font.color.rgb = title_rgb
        run.font.name = font


def build_word_document(report_md: str, graphs_folder: str, output_path: str) -> None:
    """
    Convert a markdown report to a Word .docx file.

    Handles:
      - # / ## / ### headings → Word heading styles (Deloitte green, Arial)
      - | markdown tables | → Word tables (dark-green header, white font, Arial)
      - [GRAPH: filename.png] references → embedded PNG images
      - ![alt](path) markdown images → embedded images
      - **bold** text in paragraphs
      - Regular paragraphs
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt
    except ImportError:
        logger.warning(
            "\n[WARNING] python-docx is not installed. Cannot create Word document.\n"
            "Install it with:  pip install python-docx\n"
            "The markdown report (final_report.md) was saved and is complete.\n"
        )
        return

    t0 = time.time()
    doc = Document()
    graphs_embedded = 0
    graphs_missing = 0
    tables_added = 0

    # Set page margins
    for section in doc.sections:
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)

    # ── Deloitte theme: set document-level default font to Arial ─────────────
    from docx.shared import RGBColor
    _DELOITTE_TITLE_RGB = RGBColor(0x26, 0x89, 0x0D)   # #26890D green medium
    _ARIAL = "Arial"
    try:
        normal_style = doc.styles["Normal"]
        normal_style.font.name = _ARIAL
        normal_style.font.size = Pt(10)
    except Exception:
        pass

    lines = report_md.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # ── Headers ──────────────────────────────────────────────────────────
        if stripped.startswith("### "):
            _add_deloitte_heading(doc, stripped[4:], level=3,
                                  title_rgb=_DELOITTE_TITLE_RGB, font=_ARIAL)
        elif stripped.startswith("## "):
            _add_deloitte_heading(doc, stripped[3:], level=2,
                                  title_rgb=_DELOITTE_TITLE_RGB, font=_ARIAL)
        elif stripped.startswith("# "):
            _add_deloitte_heading(doc, stripped[2:], level=1,
                                  title_rgb=_DELOITTE_TITLE_RGB, font=_ARIAL)

        # ── Markdown tables ───────────────────────────────────────────────────
        elif stripped.startswith("|") and i + 1 < len(lines) and _is_separator_row(lines[i + 1]):
            header_line = stripped
            i += 2  # Skip header + separator rows
            data_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                data_lines.append(lines[i].strip())
                i += 1
            _add_markdown_table_to_doc(doc, header_line, data_lines)
            tables_added += 1
            logger.debug(f"[DocBuilder] Added table with {len(data_lines)} data rows")
            continue  # already advanced i

        # ── Graph references: [GRAPH: filename.png] ───────────────────────────
        elif "[GRAPH:" in stripped:
            match = re.search(r"\[GRAPH:\s*([\w\-\.]+\.png)\]", stripped, re.IGNORECASE)
            if match:
                filename = match.group(1)
                img_path = Path(graphs_folder) / filename
                if img_path.exists():
                    try:
                        doc.add_picture(str(img_path), width=Inches(6.0))
                        # Caption: any text on the line outside the [GRAPH:...] tag
                        caption = re.sub(r"\[GRAPH:.*?\]", "", stripped, flags=re.IGNORECASE).strip()
                        if caption:
                            p = doc.add_paragraph(caption)
                            for run in p.runs:
                                run.italic = True
                        graphs_embedded += 1
                        logger.debug(f"[DocBuilder] Embedded graph: {filename}")
                    except Exception as e:
                        doc.add_paragraph(f"[Could not embed graph {filename}: {e}]")
                        logger.warning(f"[DocBuilder] Failed to embed graph {filename}: {e}")
                else:
                    doc.add_paragraph(f"[Graph not found: {filename}]")
                    graphs_missing += 1
                    logger.warning(f"[DocBuilder] Graph not found: {img_path}")
            else:
                _add_formatted_paragraph(doc, stripped)

        # ── Standard markdown image: ![alt](path) ────────────────────────────
        elif stripped.startswith("!["):
            match = re.search(r"!\[([^\]]*)\]\(([^\)]+)\)", stripped)
            if match:
                img_path = match.group(2)
                if Path(img_path).exists():
                    try:
                        doc.add_picture(img_path, width=Inches(6.0))
                        graphs_embedded += 1
                    except Exception as e:
                        doc.add_paragraph(f"[Could not embed image: {e}]")
                        logger.warning(f"[DocBuilder] Failed to embed image {img_path}: {e}")
                else:
                    doc.add_paragraph(f"[Image not found: {img_path}]")
                    graphs_missing += 1
            else:
                _add_formatted_paragraph(doc, stripped)

        # ── Horizontal rule ───────────────────────────────────────────────────
        elif stripped.startswith("---") and stripped == "-" * len(stripped):
            doc.add_paragraph()  # Just add spacing for HR

        # ── Non-empty text / paragraph ────────────────────────────────────────
        elif stripped:
            _add_formatted_paragraph(doc, stripped)

        # ── Empty line (ignored — paragraph spacing handles visual gaps) ───────
        # else: pass

        i += 1

    doc.save(output_path)
    elapsed = time.time() - t0
    logger.info(f"  Word document saved to: {output_path}")
    logger.info(
        f"  DocBuilder stats: {tables_added} tables, "
        f"{graphs_embedded} graphs embedded, {graphs_missing} missing"
        f" | built in {elapsed:.1f}s"
    )

# ─── Main ─────────────────────────────────────────────────────────────────────

def _generate_report_legacy(
    client, model: str, entries: list, context_text: str, available_graphs: list,
) -> str:
    """Legacy single-call report generation (fallback if iterative pipeline fails)."""
    user_message = build_report_prompt(entries, context_text, available_graphs)
    logger.info(f"  [Legacy] Single-call generation ({len(user_message):,}ch prompt)")

    report_text, usage = _llm_call(
        client, model, PHASE2_SYSTEM_PROMPT, user_message,
        PHASE2_MAX_TOKENS, label="Legacy Single-Call",
    )
    return report_text


def main() -> None:
    # Load config
    config = yaml.safe_load(read_file("config.yaml"))
    model = config.get("model", "claude-sonnet-4-6")
    archive_path = config.get("archive_file", "full_archive.txt")
    context_path = config.get("active_context_file", "active_context.md")
    graphs_folder = config.get("graphs_folder", "workspace/graphs")
    debug_logging = config.get("debug_logging", False)

    extracted_code_path = "extracted_code.md"
    report_md_path = "final_report.md"
    report_docx_path = "final_report.docx"

    # ── Setup logging ─────────────────────────────────────────────────────────
    log_file = setup_logging(debug=debug_logging)

    logger.info("=" * 60)
    logger.info("Phase 2 — Final Report Generator (Iterative Pipeline)")
    logger.info(f"Archive : {archive_path}")
    logger.info(f"Context : {context_path}")
    logger.info(f"Graphs  : {graphs_folder}")
    logger.info(f"Model   : {model}")
    logger.info(f"Log     : {log_file}")
    logger.info("=" * 60)
    logger.info("")

    # Validate inputs
    if not Path(archive_path).exists():
        logger.error(f"ERROR: {archive_path} not found. Run loop.py first.")
        sys.exit(1)

    archive_text = read_file(archive_path)
    context_text = read_file(context_path) if Path(context_path).exists() else ""
    logger.debug(
        f"[Input] archive={len(archive_text):,}ch  context={len(context_text):,}ch"
    )

    # Parse successful entries
    t_parse = time.time()
    entries = parse_archive(archive_text)
    logger.info(f"Found {len(entries)} successful analyses in archive (parsed in {time.time()-t_parse:.2f}s).")

    if not entries:
        logger.error("No successful analyses to report on. Exiting.")
        sys.exit(0)

    # ── Step A: Extract code ──────────────────────────────────────────────────
    logger.info(f"\n[Step A] Extracting analysis code to {extracted_code_path}...")
    t_a = time.time()
    extract_code_to_file(entries, extracted_code_path)
    logger.info(f"[Step A] Done in {time.time()-t_a:.1f}s")

    # ── Step B: Generate markdown report (Iterative Pipeline) ─────────────────
    available_graphs = list_available_graphs(graphs_folder)
    logger.info(f"\n[Step B] Generating markdown report (iterative pipeline)...")
    logger.info(f"         Analyses: {len(entries)}  |  Available graphs: {len(available_graphs)}")
    if available_graphs:
        for _, fname in available_graphs:
            logger.info(f"           - {fname}")
    logger.info("")

    client = anthropic.Anthropic(api_key=API_KEY)
    t_b = time.time()

    # Try iterative pipeline; fall back to legacy on failure
    report_text = generate_report_iterative(
        client, model, entries, context_text, available_graphs, config,
    )

    if report_text is None:
        logger.warning("[Step B] Iterative pipeline failed, using legacy single-call...")
        report_text = _generate_report_legacy(
            client, model, entries, context_text, available_graphs,
        )

    b_elapsed = time.time() - t_b
    write_file(report_md_path, report_text)
    logger.info(f"\nMarkdown report saved to: {report_md_path} ({len(report_text):,}ch)")
    logger.info(f"[Step B] Total time: {b_elapsed:.1f}s")

    # ── Step C: Build Word document ───────────────────────────────────────────
    logger.info(f"\n[Step C] Building Word document ({report_docx_path})...")
    t_c = time.time()
    build_word_document(report_text, graphs_folder, report_docx_path)
    logger.info(f"[Step C] Done in {time.time()-t_c:.1f}s")

    logger.info("")
    logger.info("=" * 60)
    logger.info("Phase 2 complete.")
    logger.info(f"  {extracted_code_path}  — all successful analysis code")
    logger.info(f"  {report_md_path}       — full markdown report")
    logger.info(f"  {report_docx_path}     — Word document with tables and graphs")
    logger.info(f"  {log_file}  — full debug log")
    logger.info("=" * 60)

if __name__ == "__main__":
    main()
